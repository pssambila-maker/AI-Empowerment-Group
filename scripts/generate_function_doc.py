from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

GOLD = RGBColor(0xC9, 0xA8, 0x4C)
CHARCOAL = RGBColor(0x2D, 0x2D, 0x2D)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = CHARCOAL


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = GOLD


def para(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style='List Bullet')


def table_from_rows(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = str(val)
            if i == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph("")


# Title page
title = doc.add_heading('AI Empowerment Group', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI Readiness Scorecard - Functional Specification')
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Version 1.0  |  Last updated: 11 June 2026')

doc.add_page_break()

# 1. Overview
h1('1. Overview')
para(
    "The AI Readiness Scorecard is a free, self-service assessment funnel available "
    "at /assessment on the AI Empowerment Group website. Visitors (individuals or "
    "enterprises) answer 15 questions about their current AI usage and receive an "
    "instant 0-100 readiness score, three tailored insights, and an invitation to a "
    "free recurring live class. Completed assessments are captured as leads in "
    "Firestore for follow-up."
)

h2('1.1 Goals')
bullet("Generate qualified leads (individuals and enterprises) for AI Empowerment Group services.")
bullet("Provide visitors with immediate, personalised value (a score plus insights) in exchange for their contact details.")
bullet("Convert assessment completions into free class registrations, building a pipeline for paid services.")
bullet("Be fully modular so any single piece (questions, scoring, schedule, email copy) can be edited independently without touching the rest of the system.")

# 2. User journey
h1('2. User Journey')

h2('2.1 Step 0 - Gateway')
para("The visitor lands on /assessment and chooses one of two paths:")
bullet("Individual - freelancers, coaches, consultants, solopreneurs")
bullet("Enterprise - teams, businesses, organisations, corporations")
para("This selection sets branding (badge colour and text) for the rest of the flow and is stored against the lead record.")

h2('2.2 Step 1 - Email verification (Email Gate)')
para(
    "The visitor enters their email address. The system sends a Firebase "
    "passwordless sign-in link to that address by email."
)
bullet("The visitor clicks the link in their inbox, which returns them to /assessment.")
bullet("Firebase verifies the link and signs the visitor in - no password is ever created or required.")
bullet("A Resend link option is available if the email does not arrive.")
bullet(
    "If the link is opened on a different device or browser than it was requested "
    "on, the visitor is asked to re-enter their email to confirm and complete sign-in."
)

h2('2.3 Step 2 - Profile')
para("Once signed in, the visitor provides:")
bullet("Full Name (required)")
bullet("Email Address (pre-filled, read-only - taken from the verified sign-in)")
bullet("Phone Number (optional)")

h2('2.4 Step 3 - Questions (15 total)')
para("Foundation Checklist (Questions 1-10): a multi-select checklist of statements describing AI usage habits. Each checked item is worth 4 points (40 points max).")
para("Scored Questions (Questions 11-15): five single-choice questions, each with five answer options worth 0, 3, 5, 7, or 10 points (50 points max).")
para("A progress bar tracks completion across all 15 questions (0% to 100%).")

h2('2.5 Step 4 - Results')
para("On completing question 15, the visitor immediately sees:")
bullet("An animated gauge showing their score out of 100")
bullet("A status band: Needs Attention (0-39), On Track (40-69), or High Performer (70-100)")
bullet("Three insights tailored to their score band")
bullet("A Register Free button for the recurring live class")
para(
    "At this point, the completed assessment (profile, role, score, answers) is "
    "saved to the leads collection in Firestore."
)

h2('2.6 Class registration and confirmation')
para(
    "Clicking Register Free triggers a confirmation email containing the class "
    "schedule, Zoom link, and Add to Calendar options (Google Calendar link and "
    ".ics download for Outlook/Apple Calendar). The results screen updates to show "
    "the registered state."
)

h2('2.7 Restart')
para("A Retake Assessment link returns the visitor to the Gateway step, resetting their answers (their email/sign-in session is preserved).")

# 3. Scoring
h1('3. Scoring Model')
h2('3.1 Foundation Checklist (Questions 1-10)')
para("10 statements; 4 points awarded per item checked. Maximum: 40 points.")

h2('3.2 Scored Questions (Questions 11-15)')
table_from_rows([
    ["Question", "Topic", "Point scale"],
    ["11", "How often do you use AI tools in a typical week?", "0 / 3 / 5 / 7 / 10"],
    ["12", "How confident are you at writing prompts that get great results?", "0 / 3 / 5 / 7 / 10"],
    ["13", "How much time does AI currently save you in an average week?", "0 / 3 / 5 / 7 / 10"],
    ["14", "How comfortable are you using AI for important work tasks?", "0 / 3 / 5 / 7 / 10"],
    ["15", "How would you rate your overall AI knowledge today?", "0 / 3 / 5 / 7 / 10"],
])

h2('3.3 Final score calculation')
para("Final score (0-100) = round( (checklist points + scored question points) / 90 x 100 )")
para("90 is the maximum possible raw points (40 from the checklist plus 50 from the scored questions).")

h2('3.4 Score bands')
table_from_rows([
    ["Band", "Score range", "Label shown to user"],
    ["needs-attention", "0 - 39", "Needs Attention"],
    ["on-track", "40 - 69", "On Track"],
    ["high-performer", "70 - 100", "High Performer"],
])
para("Each band has three pre-written insight statements shown on the results screen.")

# 4. Recurring free class
h1('4. Recurring Free Class')
table_from_rows([
    ["Attribute", "Value"],
    ["Day", "Saturday"],
    ["Time", "9:00 AM - 11:00 AM Eastern Time (DST-aware)"],
    ["Location", "Zoom"],
    ["Join link", "https://zoom.us/s/9832093373#success"],
])
para(
    "The system automatically calculates the date of the next upcoming Saturday "
    "occurrence (accounting for US Daylight Saving Time transitions), and uses it "
    "to populate the schedule shown on the results screen, the confirmation email, "
    "the Google Calendar link, and the downloadable .ics calendar file."
)

# 5. Email notifications
h1('5. Email Notifications')
h2('5.1 Sign-in link email')
para("Sent automatically by Firebase Authentication when a visitor requests access in Step 1. Contains a one-time secure link back to the assessment.")

h2('5.2 Class invite email')
para("Sent when the visitor clicks Register Free on the results screen. Subject: Your AI Readiness Score + Your Free Class Invite. Contains:")
bullet("A thank-you message acknowledging assessment completion")
bullet("The visitor's score")
bullet("The next class date, time, and Zoom join link")
bullet("A note encouraging the visitor to add the class to their calendar")
para("Delivery is handled by the Firebase Trigger Email extension via Gmail SMTP, triggered by writing a document to the mail Firestore collection.")

# 6. Data captured
h1('6. Data Captured (Firestore)')
h2('6.1 leads collection')
para("One document per completed assessment. Fields:")
table_from_rows([
    ["Field", "Description"],
    ["uid", "Firebase Auth user ID of the signed-in visitor"],
    ["email", "Verified email address"],
    ["fullName", "Name entered in Step 2"],
    ["phone", "Optional phone number"],
    ["role", "individual or enterprise"],
    ["score", "Final score, 0-100"],
    ["category", "needs-attention / on-track / high-performer"],
    ["checkedIndexes", "Indices of checked foundation-checklist items"],
    ["choiceAnswers", "Map of question number to selected option index"],
    ["createdAt", "Server timestamp"],
])

h2('6.2 mail collection')
para("One document per outgoing email, consumed and processed by the Trigger Email extension. Contains the recipient address and the email subject/text/HTML body.")

h2('6.3 Access and security')
para("Both collections are write-only from the client (create only) and restricted to the signed-in user - no visitor can read, edit, or delete lead or email records. All other Firestore paths are denied by default.")

# 7. System architecture
h1('7. System Architecture (Modular Design)')
para(
    "The assessment was built so that each concern lives in its own file. This "
    "means content, scoring rules, the class schedule, email copy, or any single "
    "UI step can be changed independently without affecting the rest of the system."
)

table_from_rows([
    ["File / Folder", "Responsibility"],
    ["src/config/assessment.ts", "Single source of truth: question text, scoring weights, score-band labels and insight copy, and the recurring class schedule (day, time, timezone, Zoom link)."],
    ["src/lib/assessment/scoring.ts", "Pure scoring calculation - turns checklist and answer selections into a 0-100 score and band."],
    ["src/lib/assessment/calendar.ts", "Computes the next class date/time (DST-aware), and generates Google Calendar links and .ics files."],
    ["src/lib/assessment/emailTemplates.ts", "Builds the subject/text/HTML for the class invite email."],
    ["src/lib/assessment/auth.ts", "Firebase email-link (passwordless) sign-in: sending links and completing sign-in on return."],
    ["src/lib/assessment/leads.ts", "Writes completed assessment results to the leads Firestore collection."],
    ["src/lib/assessment/mail.ts", "Queues the class invite email by writing to the mail Firestore collection."],
    ["src/lib/assessment/controller.ts", "Orchestrator - wires the UI components together into the step-by-step flow and calls the modules above."],
    ["src/components/assessment/*.astro", "One component per funnel step: Gateway, EmailGate, ProfileForm, QuestionFlow, ResultsPanel."],
    ["src/pages/assessment.astro", "Assembles the page: site layout plus the five step components plus the controller script."],
    ["src/lib/firebase/client.ts", "Shared Firebase app/auth/Firestore initialisation used across the site."],
    ["firestore.rules", "Security rules for the leads and mail collections."],
])

h2('7.1 How to make common edits')
table_from_rows([
    ["To change...", "Edit this file"],
    ["Question wording, checklist items, or point values", "src/config/assessment.ts"],
    ["Score band thresholds or insight text", "src/config/assessment.ts"],
    ["Class day/time/timezone or Zoom link", "src/config/assessment.ts (CLASS_SCHEDULE)"],
    ["Class invite email wording", "src/lib/assessment/emailTemplates.ts"],
    ["Look and feel of a single funnel step", "the matching file in src/components/assessment/"],
    ["Step ordering or flow logic", "src/lib/assessment/controller.ts"],
    ["Which Firestore collections are written to / security rules", "firestore.rules and src/lib/assessment/leads.ts / mail.ts"],
])

# 8. External dependencies
h1('8. External Services and Configuration')
bullet("Firebase Authentication - Email Link (passwordless) sign-in must remain enabled under Authentication > Sign-in method.")
bullet("Cloud Firestore - leads and mail collections, governed by firestore.rules.")
bullet("Firebase Extensions: Trigger Email (firestore-send-email) - watches the mail collection and sends emails via Gmail SMTP (configured with a Gmail App Password).")
bullet("Firebase Hosting - serves the built static site, including /assessment.")

# 9. Known issues / open items
h1('9. Known Issues and Open Items')
bullet("Sign-in link request currently returns a generic error in some environments - root cause to be confirmed (likely an Authentication configuration setting such as the Email Link toggle or Authorized domains list).")
bullet("Zoom join link is an owner-provided link and should be reviewed periodically to ensure it remains valid for the recurring Saturday class.")
bullet("Email deliverability depends on the Gmail account's sending limits (around 500/day) - monitor if lead volume grows significantly.")

doc.save('function.docx')
print("done")
